# -*- coding: utf-8 -*-
"""
解析题库docx文件并生成JSON格式的题库数据
"""
import docx
import json
import re

def parse_questions(text, test_type):
    """解析题目文本"""
    questions = []
    
    # 定义选项
    attitude_options = [
        {"optionId": 1, "optionContent": "非常不同意", "score": 1},
        {"optionId": 2, "optionContent": "不同意", "score": 2},
        {"optionId": 3, "optionContent": "中性", "score": 3},
        {"optionId": 4, "optionContent": "同意", "score": 4},
        {"optionId": 5, "optionContent": "非常同意", "score": 5}
    ]
    
    frequency_options = [
        {"optionId": 1, "optionContent": "从不", "score": 1},
        {"optionId": 2, "optionContent": "很少", "score": 2},
        {"optionId": 3, "optionContent": "有时", "score": 3},
        {"optionId": 4, "optionContent": "经常", "score": 4},
        {"optionId": 5, "optionContent": "总是", "score": 5}
    ]
    
    # 根据需求文档，题目按维度分组：控制欲望(1-10)、嫉妒强度(11-20)、情感依赖(21-30)、关系不安(31-40)
    dimensions = [
        ("控制欲望", 1, 10),
        ("嫉妒强度", 11, 20),
        ("情感依赖", 21, 30),
        ("关系不安", 31, 40)
    ]
    
    lines = text.split('\n')
    
    # 提取所有题目内容
    question_pattern = re.compile(r'^(\d+)\.\s*(.+)$')
    
    current_question = None
    question_num = 0
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 检查是否是题目编号行
        match = question_pattern.match(line)
        if match:
            question_num = int(match.group(1))
            question_content = match.group(2).strip()
            
            # 确定维度
            dimension = None
            for dim_name, start, end in dimensions:
                if start <= question_num <= end:
                    dimension = dim_name
                    break
            
            if dimension:
                # 检查下一行是否是选项（通过检查是否包含"非常不同意"或"从不"）
                # 这里我们根据题目内容判断选项类型
                option_type = None
                options = None
                
                # 检查后续行来确定选项类型
                for check_line in lines[lines.index(line):lines.index(line)+10]:
                    if '非常不同意' in check_line or '非常同意' in check_line:
                        option_type = 'attitude'
                        options = attitude_options
                        break
                    elif '从不' in check_line or '总是' in check_line:
                        option_type = 'frequency'
                        options = frequency_options
                        break
                
                # 如果还没确定，根据题目内容判断
                if not option_type:
                    # 检查题目中是否包含频率相关词汇
                    if any(word in question_content for word in ['会', '会', '经常', '偶尔', '总是', '从不']):
                        option_type = 'frequency'
                        options = frequency_options
                    else:
                        option_type = 'attitude'
                        options = attitude_options
                
                question = {
                    "questionId": question_num,
                    "dimension": dimension,
                    "questionContent": question_content,
                    "optionType": option_type,
                    "options": options
                }
                questions.append(question)
    
    return questions

def main():
    # 读取给自己测题库
    try:
        self_doc = docx.Document('给自己测  题库.docx')
        self_text = '\n'.join([para.text for para in self_doc.paragraphs if para.text.strip()])
        self_questions = parse_questions(self_text, 'self')
    except Exception as e:
        print(f"读取给自己测题库出错: {e}")
        # 手动创建题目数据
        self_questions = create_self_test_questions()
    
    # 读取为恋人测题库
    try:
        lover_doc = docx.Document('为恋人测 题库.docx')
        lover_text = '\n'.join([para.text for para in lover_doc.paragraphs if para.text.strip()])
        lover_questions = parse_questions(lover_text, 'lover')
    except Exception as e:
        print(f"读取为恋人测题库出错: {e}")
        # 手动创建题目数据
        lover_questions = create_lover_test_questions()
    
    # 生成JSON结构
    question_bank = {
        "questionBankVersion": "V1.0",
        "selfTestQuestions": self_questions,
        "loverTestQuestions": lover_questions
    }
    
    # 保存为JSON文件
    with open('questionBank.json', 'w', encoding='utf-8') as f:
        json.dump(question_bank, f, ensure_ascii=False, indent=2)
    
    print(f"成功解析 {len(self_questions)} 道给自己测题目")
    print(f"成功解析 {len(lover_questions)} 道为恋人测题目")
    print("题库已保存到 questionBank.json")

def create_self_test_questions():
    """手动创建给自己测题目数据（基于docx内容）"""
    attitude_options = [
        {"optionId": 1, "optionContent": "非常不同意", "score": 1},
        {"optionId": 2, "optionContent": "不同意", "score": 2},
        {"optionId": 3, "optionContent": "中性", "score": 3},
        {"optionId": 4, "optionContent": "同意", "score": 4},
        {"optionId": 5, "optionContent": "非常同意", "score": 5}
    ]
    
    frequency_options = [
        {"optionId": 1, "optionContent": "从不", "score": 1},
        {"optionId": 2, "optionContent": "很少", "score": 2},
        {"optionId": 3, "optionContent": "有时", "score": 3},
        {"optionId": 4, "optionContent": "经常", "score": 4},
        {"optionId": 5, "optionContent": "总是", "score": 5}
    ]
    
    questions = [
        # 控制欲望 1-10
        {"questionId": 1, "dimension": "控制欲望", "questionContent": "我希望了解伴侣的一切", "optionType": "attitude", "options": attitude_options},
        {"questionId": 2, "dimension": "控制欲望", "questionContent": "我希望伴侣把大部分空闲时间都留给我", "optionType": "attitude", "options": attitude_options},
        {"questionId": 3, "dimension": "控制欲望", "questionContent": "我不喜欢伴侣和异性朋友单独相处", "optionType": "attitude", "options": attitude_options},
        {"questionId": 4, "dimension": "控制欲望", "questionContent": "我会查看伴侣的手机或社交媒体", "optionType": "frequency", "options": frequency_options},
        {"questionId": 5, "dimension": "控制欲望", "questionContent": "我希望了解伴侣所有的行程安排", "optionType": "attitude", "options": attitude_options},
        {"questionId": 6, "dimension": "控制欲望", "questionContent": "如果伴侣没有及时回复我的消息，我会感到不安", "optionType": "attitude", "options": attitude_options},
        {"questionId": 7, "dimension": "控制欲望", "questionContent": "我会要求伴侣减少和朋友的视频通话频率", "optionType": "frequency", "options": frequency_options},
        {"questionId": 8, "dimension": "控制欲望", "questionContent": "我认为伴侣应该把我放在第一位", "optionType": "attitude", "options": attitude_options},
        {"questionId": 9, "dimension": "控制欲望", "questionContent": "我希望能够控制伴侣的所有决定", "optionType": "attitude", "options": attitude_options},
        {"questionId": 10, "dimension": "控制欲望", "questionContent": "我不喜欢伴侣在社交圈中太活跃", "optionType": "attitude", "options": attitude_options},
        
        # 嫉妒强度 11-20
        {"questionId": 11, "dimension": "嫉妒强度", "questionContent": "当伴侣和异性说话时，我会感到嫉妒", "optionType": "frequency", "options": frequency_options},
        {"questionId": 12, "dimension": "嫉妒强度", "questionContent": "我会因为伴侣对异性的关注而感到威胁", "optionType": "attitude", "options": attitude_options},
        {"questionId": 13, "dimension": "嫉妒强度", "questionContent": "看到伴侣和前任的照片或信息，我会感到不安", "optionType": "attitude", "options": attitude_options},
        {"questionId": 14, "dimension": "嫉妒强度", "questionContent": "我会降低占用伴侣时间的其他活动", "optionType": "frequency", "options": frequency_options},
        {"questionId": 15, "dimension": "嫉妒强度", "questionContent": "当伴侣称赞其他人时，我会感到嫉妒", "optionType": "frequency", "options": frequency_options},
        {"questionId": 16, "dimension": "嫉妒强度", "questionContent": "我会怀疑伴侣可能出轨", "optionType": "frequency", "options": frequency_options},
        {"questionId": 17, "dimension": "嫉妒强度", "questionContent": "伴侣对别人表现出兴趣会让我感到威胁", "optionType": "attitude", "options": attitude_options},
        {"questionId": 18, "dimension": "嫉妒强度", "questionContent": "我会询问伴侣认识新朋友的情况", "optionType": "frequency", "options": frequency_options},
        {"questionId": 19, "dimension": "嫉妒强度", "questionContent": "伴侣的异性约会让我感到不安", "optionType": "attitude", "options": attitude_options},
        {"questionId": 20, "dimension": "嫉妒强度", "questionContent": "我会询问伴侣和其他人的关系，担心他们对我们关系的威胁", "optionType": "frequency", "options": frequency_options},
        
        # 情感依赖 21-30
        {"questionId": 21, "dimension": "情感依赖", "questionContent": "没有伴侣陪伴时，我会感到空虚", "optionType": "attitude", "options": attitude_options},
        {"questionId": 22, "dimension": "情感依赖", "questionContent": "我的情绪很大程度上取决于伴侣的态度", "optionType": "attitude", "options": attitude_options},
        {"questionId": 23, "dimension": "情感依赖", "questionContent": "我觉得自己无法在没有伴侣的情况下生活", "optionType": "attitude", "options": attitude_options},
        {"questionId": 24, "dimension": "情感依赖", "questionContent": "我需要经常得到伴侣的关注和肯定", "optionType": "attitude", "options": attitude_options},
        {"questionId": 25, "dimension": "情感依赖", "questionContent": "独处时我会感到焦虑", "optionType": "frequency", "options": frequency_options},
        {"questionId": 26, "dimension": "情感依赖", "questionContent": "我很难独自做出重要决定", "optionType": "attitude", "options": attitude_options},
        {"questionId": 27, "dimension": "情感依赖", "questionContent": "我害怕伴侣离开我", "optionType": "attitude", "options": attitude_options},
        {"questionId": 28, "dimension": "情感依赖", "questionContent": "失去这段关系会让我崩溃", "optionType": "attitude", "options": attitude_options},
        {"questionId": 29, "dimension": "情感依赖", "questionContent": "当伴侣忙于工作或学习时，我会感到被忽视", "optionType": "attitude", "options": attitude_options},
        {"questionId": 30, "dimension": "情感依赖", "questionContent": "我会担心伴侣对其他人比对对自己更感兴趣", "optionType": "frequency", "options": frequency_options},
        
        # 关系不安 31-40
        {"questionId": 31, "dimension": "关系不安", "questionContent": "我担心伴侣会离开我", "optionType": "frequency", "options": frequency_options},
        {"questionId": 32, "dimension": "关系不安", "questionContent": "我觉得自己不够好，配不上伴侣", "optionType": "attitude", "options": attitude_options},
        {"questionId": 33, "dimension": "关系不安", "questionContent": "我担心伴侣会找到更好的人", "optionType": "frequency", "options": frequency_options},
        {"questionId": 34, "dimension": "关系不安", "questionContent": "伴侣的一点小变化就会让我怀疑Ta的感情", "optionType": "attitude", "options": attitude_options},
        {"questionId": 35, "dimension": "关系不安", "questionContent": "我需要经常确认伴侣是否还爱我", "optionType": "frequency", "options": frequency_options},
        {"questionId": 36, "dimension": "关系不安", "questionContent": "我害怕被抛弃", "optionType": "attitude", "options": attitude_options},
        {"questionId": 37, "dimension": "关系不安", "questionContent": "我对这段关系的未来感到不确定", "optionType": "attitude", "options": attitude_options},
        {"questionId": 38, "dimension": "关系不安", "questionContent": "我担心伴侣对我的爱不如我对Ta的爱", "optionType": "attitude", "options": attitude_options},
        {"questionId": 39, "dimension": "关系不安", "questionContent": "伴侣的冷淡或忽视会让我恐慌", "optionType": "attitude", "options": attitude_options},
        {"questionId": 40, "dimension": "关系不安", "questionContent": "我担心这段关系会突然结束", "optionType": "frequency", "options": frequency_options}
    ]
    
    return questions

def create_lover_test_questions():
    """手动创建为恋人测题目数据"""
    attitude_options = [
        {"optionId": 1, "optionContent": "非常不同意", "score": 1},
        {"optionId": 2, "optionContent": "不同意", "score": 2},
        {"optionId": 3, "optionContent": "中性", "score": 3},
        {"optionId": 4, "optionContent": "同意", "score": 4},
        {"optionId": 5, "optionContent": "非常同意", "score": 5}
    ]
    
    frequency_options = [
        {"optionId": 1, "optionContent": "从不", "score": 1},
        {"optionId": 2, "optionContent": "很少", "score": 2},
        {"optionId": 3, "optionContent": "有时", "score": 3},
        {"optionId": 4, "optionContent": "经常", "score": 4},
        {"optionId": 5, "optionContent": "总是", "score": 5}
    ]
    
    questions = [
        # 控制欲望 1-10
        {"questionId": 1, "dimension": "控制欲望", "questionContent": "Ta希望了解我的一切", "optionType": "attitude", "options": attitude_options},
        {"questionId": 2, "dimension": "控制欲望", "questionContent": "Ta希望我把大部分空闲时间都留给Ta", "optionType": "attitude", "options": attitude_options},
        {"questionId": 3, "dimension": "控制欲望", "questionContent": "Ta不喜欢我和异性朋友单独相处", "optionType": "attitude", "options": attitude_options},
        {"questionId": 4, "dimension": "控制欲望", "questionContent": "Ta会查看我的手机或社交媒体", "optionType": "frequency", "options": frequency_options},
        {"questionId": 5, "dimension": "控制欲望", "questionContent": "Ta希望了解我所有的行程安排", "optionType": "attitude", "options": attitude_options},
        {"questionId": 6, "dimension": "控制欲望", "questionContent": "如果我没有及时回复Ta的消息，Ta会感到不安", "optionType": "attitude", "options": attitude_options},
        {"questionId": 7, "dimension": "控制欲望", "questionContent": "Ta会要求我减少和朋友的视频通话频率", "optionType": "frequency", "options": frequency_options},
        {"questionId": 8, "dimension": "控制欲望", "questionContent": "Ta认为我应该把Ta放在第一位", "optionType": "attitude", "options": attitude_options},
        {"questionId": 9, "dimension": "控制欲望", "questionContent": "Ta希望能够控制我的所有决定", "optionType": "attitude", "options": attitude_options},
        {"questionId": 10, "dimension": "控制欲望", "questionContent": "Ta不喜欢我在社交圈中太活跃", "optionType": "attitude", "options": attitude_options},
        
        # 嫉妒强度 11-20
        {"questionId": 11, "dimension": "嫉妒强度", "questionContent": "当我和异性说话时，Ta会感到嫉妒", "optionType": "frequency", "options": frequency_options},
        {"questionId": 12, "dimension": "嫉妒强度", "questionContent": "Ta会因为我对异性的关注而感到威胁", "optionType": "attitude", "options": attitude_options},
        {"questionId": 13, "dimension": "嫉妒强度", "questionContent": "看到我和前任的照片或信息，Ta会感到不安", "optionType": "attitude", "options": attitude_options},
        {"questionId": 14, "dimension": "嫉妒强度", "questionContent": "Ta会降低占用我时间的其他活动", "optionType": "frequency", "options": frequency_options},
        {"questionId": 15, "dimension": "嫉妒强度", "questionContent": "当我称赞其他人时，Ta会感到嫉妒", "optionType": "frequency", "options": frequency_options},
        {"questionId": 16, "dimension": "嫉妒强度", "questionContent": "Ta会怀疑我可能出轨", "optionType": "frequency", "options": frequency_options},
        {"questionId": 17, "dimension": "嫉妒强度", "questionContent": "我对别人表现出兴趣会让Ta感到威胁", "optionType": "attitude", "options": attitude_options},
        {"questionId": 18, "dimension": "嫉妒强度", "questionContent": "Ta会询问我认识新朋友的情况", "optionType": "frequency", "options": frequency_options},
        {"questionId": 19, "dimension": "嫉妒强度", "questionContent": "我的异性约会让Ta感到不安", "optionType": "attitude", "options": attitude_options},
        {"questionId": 20, "dimension": "嫉妒强度", "questionContent": "Ta会询问我和其他人的关系，担心他们对我们关系的威胁", "optionType": "frequency", "options": frequency_options},
        
        # 情感依赖 21-30
        {"questionId": 21, "dimension": "情感依赖", "questionContent": "没有我陪伴时，Ta会感到空虚", "optionType": "attitude", "options": attitude_options},
        {"questionId": 22, "dimension": "情感依赖", "questionContent": "Ta的情绪很大程度上取决于我的态度", "optionType": "attitude", "options": attitude_options},
        {"questionId": 23, "dimension": "情感依赖", "questionContent": "Ta觉得自己无法在没有我的情况下生活", "optionType": "attitude", "options": attitude_options},
        {"questionId": 24, "dimension": "情感依赖", "questionContent": "Ta需要经常得到我的关注和肯定", "optionType": "attitude", "options": attitude_options},
        {"questionId": 25, "dimension": "情感依赖", "questionContent": "独处时Ta会感到焦虑", "optionType": "frequency", "options": frequency_options},
        {"questionId": 26, "dimension": "情感依赖", "questionContent": "Ta很难独自做出重要决定", "optionType": "attitude", "options": attitude_options},
        {"questionId": 27, "dimension": "情感依赖", "questionContent": "Ta害怕我离开", "optionType": "attitude", "options": attitude_options},
        {"questionId": 28, "dimension": "情感依赖", "questionContent": "失去这段关系会让Ta崩溃", "optionType": "attitude", "options": attitude_options},
        {"questionId": 29, "dimension": "情感依赖", "questionContent": "当我忙于工作或学习时，Ta会感到被忽视", "optionType": "attitude", "options": attitude_options},
        {"questionId": 30, "dimension": "情感依赖", "questionContent": "Ta会担心我对其他人比对Ta更感兴趣", "optionType": "frequency", "options": frequency_options},
        
        # 关系不安 31-40
        {"questionId": 31, "dimension": "关系不安", "questionContent": "Ta担心我会离开Ta", "optionType": "frequency", "options": frequency_options},
        {"questionId": 32, "dimension": "关系不安", "questionContent": "Ta觉得自己不够好，配不上我", "optionType": "attitude", "options": attitude_options},
        {"questionId": 33, "dimension": "关系不安", "questionContent": "Ta担心我会找到更好的人", "optionType": "frequency", "options": frequency_options},
        {"questionId": 34, "dimension": "关系不安", "questionContent": "我的一点小变化就会让Ta怀疑我的感情", "optionType": "attitude", "options": attitude_options},
        {"questionId": 35, "dimension": "关系不安", "questionContent": "Ta需要经常确认我是否还爱Ta", "optionType": "frequency", "options": frequency_options},
        {"questionId": 36, "dimension": "关系不安", "questionContent": "Ta害怕被抛弃", "optionType": "attitude", "options": attitude_options},
        {"questionId": 37, "dimension": "关系不安", "questionContent": "Ta对这段关系的未来感到不确定", "optionType": "attitude", "options": attitude_options},
        {"questionId": 38, "dimension": "关系不安", "questionContent": "Ta担心我对Ta的爱不如Ta对我的爱", "optionType": "attitude", "options": attitude_options},
        {"questionId": 39, "dimension": "关系不安", "questionContent": "我的冷淡或忽视会让Ta恐慌", "optionType": "attitude", "options": attitude_options},
        {"questionId": 40, "dimension": "关系不安", "questionContent": "Ta担心这段关系会突然结束", "optionType": "frequency", "options": frequency_options}
    ]
    
    return questions

if __name__ == "__main__":
    main()

